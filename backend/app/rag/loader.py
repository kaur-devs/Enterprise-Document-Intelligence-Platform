import os
from typing import List
from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_community.document_loaders import PyPDFLoader, TextLoader

HEADING_STYLE_LEVELS = {"Heading 1": 1, "Heading 2": 2, "Heading 3": 3}

class DocumentLoader:
    @staticmethod
    def _load_docx_as_markdown(filepath: str) -> List[Document]:
        docx_file = DocxDocument(filepath)
        lines = []
        for paragraph in docx_file.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            level = HEADING_STYLE_LEVELS.get(paragraph.style.name)
            lines.append(f"{'#' * level} {text}" if level else text)
        return [Document(page_content="\n\n".join(lines), metadata={"source": filepath})]

    @staticmethod
    def load_document(filepath: str, file_type: str) -> List[Document]:
        if not os.path.exists(filepath):
            raise FileNotFoundError(f"File not found: {filepath}")
        file_type = file_type.lower()
        if file_type == "pdf":
            return PyPDFLoader(filepath).load()
        elif file_type in ["docx", "doc"]:
            return DocumentLoader._load_docx_as_markdown(filepath)
        elif file_type in ["md", "markdown", "txt"]:
            return TextLoader(filepath, encoding="utf-8").load()
        else:
            raise ValueError(f"Unsupported file format: {file_type}")
